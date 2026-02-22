# -*- coding: utf-8 -*-
"""
Construtor de relatório DOCX profissional para auditoria de gás natural.
Converte Markdown gerado pelo Gemini em elementos python-docx formatados.
Suporta: equações LaTeX→OMML, sumário manual, tabelas de dados, acentuação correta.
"""
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Equações LaTeX → OMML
try:
    import latex2mathml.converter
    from lxml import etree
    HAS_EQUATION_SUPPORT = True
except ImportError:
    HAS_EQUATION_SUPPORT = False

# Cores do tema
AZUL_ESCURO = RGBColor(0x1A, 0x23, 0x7E)
AZUL_MEDIO = RGBColor(0x21, 0x96, 0xF3)
LARANJA = RGBColor(0xFF, 0x98, 0x00)
VERDE = RGBColor(0x4C, 0xAF, 0x50)
CINZA_CLARO = RGBColor(0xF5, 0xF5, 0xF5)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_TEXTO = RGBColor(0x66, 0x66, 0x66)
CINZA_HEADER = RGBColor(0x99, 0x99, 0x99)

# XSLT para converter MathML → OMML
# Tenta encontrar o XSL do Office; senão usa fallback
MML2OMML_XSL = None
_xsl_paths = [
    r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
    r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
    r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
    r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
]
for _p in _xsl_paths:
    if os.path.exists(_p):
        MML2OMML_XSL = _p
        break


class AuditReportBuilder:
    """Constrói um DOCX profissional de relatório de auditoria."""

    def __init__(self, graficos_dir: str):
        self.doc = Document()
        self.graficos_dir = graficos_dir
        self._heading_list = []  # Para sumário manual
        self._setup_styles()

    def _setup_styles(self):
        """Configura estilos base do documento."""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = AZUL_ESCURO
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)

        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = AZUL_MEDIO
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)

        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Calibri'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = AZUL_ESCURO
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)

        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

    def add_cover_page(self):
        """Página de capa profissional."""
        for _ in range(6):
            self.doc.add_paragraph("")

        # Instituição
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Agência Regulatória")
        run.font.size = Pt(16)
        run.font.color.rgb = AZUL_MEDIO
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Diretoria de Gás e Energia")
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL_ESCURO

        self.doc.add_paragraph("")

        # Linha separadora
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 60)
        run.font.color.rgb = AZUL_MEDIO

        self.doc.add_paragraph("")

        # Título
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RELATÓRIO DE AUDITORIA TÉCNICA")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = AZUL_ESCURO

        self.doc.add_paragraph("")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Análise das Condições de Operação\nde Distrito de Distribuição de Gás Natural")
        run.font.size = Pt(14)
        run.font.color.rgb = AZUL_ESCURO

        self.doc.add_paragraph("")

        # Linha separadora
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 60)
        run.font.color.rgb = AZUL_MEDIO

        for _ in range(3):
            self.doc.add_paragraph("")

        # Metadados
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Período Analisado: Abril a Setembro de 2025")
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL_ESCURO

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        data_atual = datetime.now().strftime("%d/%m/%Y")
        run = p.add_run(f"Data de Emissão: {data_atual}")
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL_ESCURO

        self.doc.add_paragraph("")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Medições Inteligentes e Gestão Integrada")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = AZUL_MEDIO

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENCIAL")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = LARANJA

    def add_table_of_contents(self, section_titles: list):
        """Sumário manual com títulos reais das seções."""
        self.doc.add_page_break()
        self.doc.add_heading("Sumário", level=1)

        self.doc.add_paragraph("")

        for title in section_titles:
            p = self.doc.add_paragraph()
            # Determinar nível pelo prefixo
            if title.startswith("   "):
                # Sub-seção (nível 3)
                run = p.add_run("    " + title.strip())
                run.font.size = Pt(10)
                run.font.color.rgb = CINZA_TEXTO
            else:
                run = p.add_run(title)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = AZUL_ESCURO
            p.paragraph_format.space_after = Pt(3)

        # Nota sobre atualização
        self.doc.add_paragraph("")
        p = self.doc.add_paragraph()
        run = p.add_run("Nota: Para números de página atualizados, use Ctrl+A → F9 no Word.")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = CINZA_HEADER

        # Também inserir campo TOC como fallback
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        p2 = self.doc.add_paragraph()
        r = p2.add_run()
        r._r.append(fldChar1)
        r2 = p2.add_run()
        r2._r.append(instrText)
        r3 = p2.add_run()
        r3._r.append(fldChar2)
        r4 = p2.add_run()
        r4._r.append(fldChar3)

    def add_headers_footers(self):
        """Cabeçalhos e rodapés com paginação."""
        section = self.doc.sections[0]

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = "Agência Regulatória — Relatório de Auditoria Técnica — Distrito de Gás Natural"
        hp.style.font.size = Pt(8)
        hp.style.font.color.rgb = CINZA_HEADER
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = fp.add_run("Página ")
        run1.font.size = Pt(8)
        run1.font.color.rgb = CINZA_HEADER

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

        run2 = fp.add_run()
        run2._r.append(fldChar1)
        run3 = fp.add_run()
        run3._r.append(instrText)
        run4 = fp.add_run()
        run4._r.append(fldChar2)
        run5 = fp.add_run()
        run5._r.append(fldChar3)

        run6 = fp.add_run(" | Confidencial")
        run6.font.size = Pt(8)
        run6.font.color.rgb = CINZA_HEADER

    def add_data_table(self, titulo: str, headers: list, rows: list):
        """Insere tabela de dados estatísticos formatada."""
        # Título da tabela
        p = self.doc.add_paragraph()
        run = p.add_run(titulo)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = AZUL_ESCURO
        p.paragraph_format.space_after = Pt(4)

        n_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, header_text in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = BRANCO
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A237E"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Rows
        for i, row_data in enumerate(rows):
            for j in range(min(len(row_data), n_cols)):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(row_data[j])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            if i % 2 == 0:
                for j in range(n_cols):
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    table.rows[i + 1].cells[j]._tc.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph("")

    def add_equation(self, latex_str: str, display: bool = True):
        """Converte LaTeX para equação Word nativa (OMML)."""
        if not HAS_EQUATION_SUPPORT:
            # Fallback: texto Unicode em Cambria Math
            p = self.doc.add_paragraph()
            if display:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self._latex_to_unicode(latex_str))
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.font.italic = True
            return p

        try:
            # LaTeX → MathML
            mathml_str = latex2mathml.converter.convert(latex_str)

            if MML2OMML_XSL:
                # MathML → OMML via XSLT
                mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
                xslt_tree = etree.parse(MML2OMML_XSL)
                transform = etree.XSLT(xslt_tree)
                omml_tree = transform(mathml_tree)

                p = self.doc.add_paragraph()
                if display:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p._element.append(omml_tree.getroot())
                return p
            else:
                # Sem XSL: fallback para texto Unicode em Cambria Math
                p = self.doc.add_paragraph()
                if display:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(self._latex_to_unicode(latex_str))
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
                run.font.italic = True
                return p

        except Exception:
            # Qualquer erro: fallback Unicode seguro
            p = self.doc.add_paragraph()
            if display:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self._latex_to_unicode(latex_str))
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.font.italic = True
            return p

    def add_section_from_markdown(self, markdown_text: str):
        """Converte Markdown do Gemini em elementos DOCX."""
        lines = markdown_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Linhas vazias
            if not stripped:
                i += 1
                continue

            # Display equations: $$...$$
            if stripped.startswith('$$'):
                eq_lines = [stripped[2:]]
                if stripped.endswith('$$') and len(stripped) > 4:
                    # Single-line equation
                    latex = stripped[2:-2].strip()
                    self.add_equation(latex, display=True)
                    i += 1
                    continue
                else:
                    # Multi-line equation
                    i += 1
                    while i < len(lines):
                        l = lines[i].strip()
                        if l.endswith('$$'):
                            eq_lines.append(l[:-2])
                            break
                        eq_lines.append(l)
                        i += 1
                    latex = ' '.join(eq_lines).strip()
                    self.add_equation(latex, display=True)
                    i += 1
                    continue

            # Headings
            if stripped.startswith('#### '):
                self.doc.add_heading(stripped[5:].strip(), level=4)
                i += 1
                continue
            if stripped.startswith('### '):
                self.doc.add_heading(stripped[4:].strip(), level=3)
                i += 1
                continue
            if stripped.startswith('## '):
                title = stripped[3:].strip()
                # Pular headings de seção duplicados (já adicionados pelo orquestrador)
                if title.startswith("SEÇÃO") or title.startswith("SECAO"):
                    i += 1
                    continue
                self.doc.add_heading(title, level=2)
                i += 1
                continue

            # Separadores
            if stripped in ('---', '***', '___'):
                i += 1
                continue

            # Tabelas Markdown
            if '|' in stripped and stripped.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_markdown_table(table_lines)
                continue

            # Bullet lists
            if stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # Numbered lists
            match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if match:
                text = match.group(2)
                p = self.doc.add_paragraph(style='List Number')
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # Regular paragraph
            p = self.doc.add_paragraph()
            self._add_formatted_runs(p, stripped)
            i += 1

    @staticmethod
    def _latex_to_unicode(latex: str) -> str:
        """Converte comandos LaTeX comuns para Unicode legível."""
        s = latex
        # Remover \text{...} → conteúdo
        s = re.sub(r'\\text\s*\{([^}]*)\}', r'\1', s)
        # Remover \mathrm{...} → conteúdo
        s = re.sub(r'\\mathrm\s*\{([^}]*)\}', r'\1', s)
        # Frações simples \frac{a}{b} → a/b
        s = re.sub(r'\\frac\s*\{([^}]*)\}\s*\{([^}]*)\}', r'(\1)/(\2)', s)
        # Raiz quadrada \sqrt{...} → √(...)
        s = re.sub(r'\\sqrt\s*\{([^}]*)\}', r'√(\1)', s)
        # Somatório
        s = s.replace('\\sum', 'Σ')
        # Operadores e símbolos
        s = s.replace('\\pm', '±')
        s = s.replace('\\mp', '∓')
        s = s.replace('\\times', '×')
        s = s.replace('\\cdot', '·')
        s = s.replace('\\div', '÷')
        s = s.replace('\\approx', '≈')
        s = s.replace('\\neq', '≠')
        s = s.replace('\\leq', '≤')
        s = s.replace('\\geq', '≥')
        s = s.replace('\\ll', '≪')
        s = s.replace('\\gg', '≫')
        s = s.replace('\\infty', '∞')
        s = s.replace('\\Delta', 'Δ')
        s = s.replace('\\delta', 'δ')
        s = s.replace('\\sigma', 'σ')
        s = s.replace('\\mu', 'μ')
        s = s.replace('\\alpha', 'α')
        s = s.replace('\\beta', 'β')
        s = s.replace('\\gamma', 'γ')
        s = s.replace('\\lambda', 'λ')
        s = s.replace('\\pi', 'π')
        s = s.replace('\\rho', 'ρ')
        # Símbolos escapados
        s = s.replace('\\%', '%')
        s = s.replace('\\#', '#')
        s = s.replace('\\&', '&')
        s = s.replace('\\$', '$')
        s = s.replace('\\left', '')
        s = s.replace('\\right', '')
        # Espaços LaTeX
        s = s.replace('\\quad', '  ')
        s = s.replace('\\qquad', '    ')
        s = s.replace('\\,', ' ')
        s = s.replace('\\;', ' ')
        s = s.replace('\\:', ' ')
        s = s.replace('\\ ', ' ')
        # Superscripts comuns
        s = s.replace('^2', '²')
        s = s.replace('^3', '³')
        s = s.replace('^{2}', '²')
        s = s.replace('^{3}', '³')
        # Remover chaves restantes de agrupamento
        s = s.replace('{', '')
        s = s.replace('}', '')
        # Remover backslashes restantes de comandos não reconhecidos
        s = re.sub(r'\\([a-zA-Z]+)', r'\1', s)
        # Limpar espaços múltiplos
        s = re.sub(r'  +', ' ', s)
        return s.strip()

    def _add_inline_equation(self, paragraph, latex: str):
        """Tenta inserir equação inline como OMML; fallback para Unicode + Cambria Math."""
        if HAS_EQUATION_SUPPORT and MML2OMML_XSL:
            try:
                mathml_str = latex2mathml.converter.convert(latex)
                mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
                xslt_tree = etree.parse(MML2OMML_XSL)
                transform = etree.XSLT(xslt_tree)
                omml_tree = transform(mathml_tree)
                # Inserir OMML inline no parágrafo
                omml_root = omml_tree.getroot()
                paragraph._element.append(omml_root)
                return
            except Exception:
                pass  # Fallback abaixo

        # Fallback: converter LaTeX para Unicode legível
        text = self._latex_to_unicode(latex)
        run = paragraph.add_run(text)
        run.font.name = 'Cambria Math'
        run.font.italic = True

    def _add_formatted_runs(self, paragraph, text: str):
        """Adiciona runs com formatação bold/italic e equações inline."""
        # Split por padrões: **bold**, *italic*, ***bold-italic***, $inline eq$
        pattern = r'(\*\*\*[^*]+?\*\*\*|\*\*[^*]+?\*\*|\*[^*]+?\*|\$[^$]+?\$)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('***') and part.endswith('***') and len(part) > 6:
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('$') and part.endswith('$') and len(part) > 2:
                # Inline equation — tenta OMML, fallback Unicode
                latex = part[1:-1]
                self._add_inline_equation(paragraph, latex)
            else:
                paragraph.add_run(part)

    def _add_markdown_table(self, table_lines: list):
        """Converte tabela Markdown em tabela DOCX formatada."""
        if len(table_lines) < 2:
            return

        headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]

        data_start = 1
        if len(table_lines) > 1 and re.match(r'^[\|\s\-:]+$', table_lines[1]):
            data_start = 2

        rows = []
        for line in table_lines[data_start:]:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                rows.append(cells)

        if not headers:
            return

        n_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, header_text in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = BRANCO
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A237E"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        for i, row_data in enumerate(rows):
            for j in range(min(len(row_data), n_cols)):
                cell = table.rows[i + 1].cells[j]
                # Limpar formatação markdown das células
                cell_text = row_data[j]
                cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            if i % 2 == 0:
                for j in range(n_cols):
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    table.rows[i + 1].cells[j]._tc.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph("")

    def add_graph(self, filename: str, caption: str, width_inches: float = 5.5):
        """Insere gráfico PNG com legenda centralizada."""
        path = os.path.join(self.graficos_dir, filename)
        if os.path.exists(path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width_inches))

            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = CINZA_TEXTO

            self.doc.add_paragraph("")
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Gráfico não encontrado: {filename}]")
            run.font.italic = True
            run.font.color.rgb = LARANJA

    def add_diagram(self, filepath: str, caption: str, width_inches: float = 6.0):
        """Insere diagrama de processo (imagem de diagramas/)."""
        if os.path.exists(filepath):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(filepath, width=Inches(width_inches))

            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = CINZA_TEXTO
            self.doc.add_paragraph("")

    def add_chapter_structured(
        self,
        title: str,
        introducao: str,
        tabela: dict | None,
        metodologia_text: str,
        dados_text: str,
        graph_items: list,
        graficos_text: str,
        parecer_text: str,
        diagram_items: list | None = None,
    ):
        """
        Monta capítulo completo na ordem correta:
        1. Título (Heading 1)
        2. Introdução (parágrafo da síntese)
        3. Diagramas (se houver, apenas Cap 1)
        4. Tabela de dados (se houver)
        5. ### Fundamentação Teórica
        6. ### Análise dos Dados
        7. Gráficos com legendas (ANTES da discussão)
        8. ### Discussão dos Gráficos
        9. ### Parecer Regulatório

        Args:
            title: Título do capítulo (ex: "2. Análise de Volumes de Entrada")
            introducao: Texto introdutório (sem heading)
            tabela: dict com titulo, headers, rows (ou None)
            metodologia_text: Texto da Fundamentação Teórica
            dados_text: Texto da Análise dos Dados
            graph_items: Lista de (filename, caption) dos gráficos
            graficos_text: Texto da Discussão dos Gráficos
            parecer_text: Texto do Parecer Regulatório
            diagram_items: Lista de (filepath, caption) dos diagramas (Cap 1)
        """
        self.doc.add_page_break()
        self.doc.add_heading(title, level=1)

        # 1. Introdução
        if introducao and introducao.strip():
            self.add_section_from_markdown(introducao)

        # 2. Diagramas (apenas Cap 1)
        if diagram_items:
            for filepath, caption in diagram_items:
                self.add_diagram(filepath, caption)

        # 3. Tabela de dados
        if tabela:
            self.add_data_table(tabela["titulo"], tabela["headers"], tabela["rows"])

        # 4. Fundamentação Teórica
        if metodologia_text and metodologia_text.strip():
            self.doc.add_heading("Fundamentação Teórica", level=2)
            self._add_subsection_content(metodologia_text)

        # 5. Análise dos Dados
        if dados_text and dados_text.strip():
            self.doc.add_heading("Análise dos Dados", level=2)
            self._add_subsection_content(dados_text)

        # 6. Gráficos (ANTES da discussão)
        if graph_items:
            for graph_file, caption in graph_items:
                self.add_graph(graph_file, caption)

        # 7. Discussão dos Gráficos
        if graficos_text and graficos_text.strip():
            self.doc.add_heading("Discussão dos Gráficos", level=2)
            self._add_subsection_content(graficos_text)

        # 8. Parecer Regulatório
        if parecer_text and parecer_text.strip():
            self.doc.add_heading("Parecer Regulatório", level=2)
            self._add_subsection_content(parecer_text)

    def _add_subsection_content(self, text: str):
        """Adiciona texto de subseção, removendo headings duplicados no topo."""
        # Remove headings que repetem o nome da subseção (já adicionado pelo caller)
        lines = text.split("\n")
        cleaned_lines = []
        skip_first_heading = True
        for line in lines:
            stripped = line.strip()
            if skip_first_heading and stripped.startswith("#"):
                # Pula qualquer heading no início (## ou ###)
                skip_first_heading = False
                continue
            skip_first_heading = False
            cleaned_lines.append(line)
        cleaned = "\n".join(cleaned_lines).strip()
        if cleaned:
            self.add_section_from_markdown(cleaned)

    # =================================================================
    # APÊNDICE: NOTEBOOKS JUPYTER
    # =================================================================

    MAX_OUTPUT_LINES = 80

    def _apply_code_style(self, paragraph, bg_color="F5F5F5", border_color="1A3C6E"):
        """Aplica fundo e borda esquerda a um parágrafo de código/output."""
        pPr = paragraph._element.get_or_add_pPr()
        # Fundo
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>'
        )
        pPr.append(shading)
        # Borda esquerda
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="12" w:space="4" w:color="{border_color}"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)

    def add_code_cell(self, source_lines: list, execution_count=None):
        """Adiciona uma célula de código Python formatada."""
        code = "".join(source_lines).rstrip()
        if not code:
            return

        # Label "In [N]:"
        label = f"In [{execution_count or ' '}]:"
        p_label = self.doc.add_paragraph()
        run = p_label.add_run(label)
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        p_label.paragraph_format.space_after = Pt(2)
        p_label.paragraph_format.space_before = Pt(8)

        # Linhas de código
        for line in code.split("\n"):
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.3)
            self._apply_code_style(p, bg_color="F5F5F5", border_color="1A3C6E")

    def add_output_cell(self, outputs: list):
        """Adiciona os outputs de uma célula (stream, execute_result, display_data)."""
        import base64
        import io

        for output in outputs:
            output_type = output.get("output_type", "")

            # --- Imagem (display_data com image/png) ---
            if output_type == "display_data" and "image/png" in output.get("data", {}):
                try:
                    img_b64 = output["data"]["image/png"]
                    # Pode ser string ou lista
                    if isinstance(img_b64, list):
                        img_b64 = "".join(img_b64)
                    img_data = base64.b64decode(img_b64)
                    img_stream = io.BytesIO(img_data)
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(5.0))
                except Exception:
                    pass  # Silently skip broken images
                continue

            # --- Texto (stream ou execute_result) ---
            text_lines = []
            if output_type == "stream":
                raw = output.get("text", [])
                if isinstance(raw, list):
                    text_lines = "".join(raw).split("\n")
                else:
                    text_lines = str(raw).split("\n")
            elif output_type == "execute_result":
                data = output.get("data", {})
                raw = data.get("text/plain", "")
                if isinstance(raw, list):
                    text_lines = "".join(raw).split("\n")
                else:
                    text_lines = str(raw).split("\n")

            if not text_lines:
                continue

            # Truncar se muito longo
            truncated = False
            if len(text_lines) > self.MAX_OUTPUT_LINES:
                omitted = len(text_lines) - self.MAX_OUTPUT_LINES
                text_lines = text_lines[:self.MAX_OUTPUT_LINES]
                truncated = True

            # Label "Out:"
            p_label = self.doc.add_paragraph()
            run = p_label.add_run("Out:")
            run.bold = True
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x2D, 0x50, 0x16)
            p_label.paragraph_format.space_after = Pt(2)
            p_label.paragraph_format.space_before = Pt(4)

            for line in text_lines:
                p = self.doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x2D, 0x50, 0x16)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Inches(0.3)
                self._apply_code_style(p, bg_color="F0F8F0", border_color="2D5016")

            if truncated:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[... saída truncada: {omitted} linhas omitidas]")
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True
                p.paragraph_format.left_indent = Inches(0.3)

    def add_notebook_appendix(self, notebooks: list):
        """
        Adiciona apêndice com páginas dos notebooks (PDF → imagens).
        notebooks: lista de {"path": str (PDF), "titulo": str}
        """
        import fitz  # PyMuPDF
        from io import BytesIO

        self.doc.add_page_break()
        self.doc.add_heading(
            "Apêndice A — Código-Fonte e Resultados dos Notebooks", level=1
        )

        intro = self.doc.add_paragraph()
        run = intro.add_run(
            "Este apêndice apresenta o código-fonte Python e os respectivos "
            "resultados de execução dos 7 notebooks Jupyter utilizados na "
            "análise de dados desta auditoria. Os notebooks foram executados "
            "no Google Colab e seus outputs preservados integralmente."
        )
        run.font.size = Pt(11)
        intro.paragraph_format.space_after = Pt(12)

        for idx, nb_info in enumerate(notebooks, 1):
            pdf_path = nb_info["path"]
            titulo = nb_info["titulo"]

            if idx > 1:
                self.doc.add_page_break()

            self.doc.add_heading(f"A.{idx} {titulo}", level=2)

            try:
                pdf_doc = fitz.open(pdf_path)
            except Exception as e:
                p = self.doc.add_paragraph()
                run = p.add_run(f"Erro ao carregar PDF: {e}")
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                continue

            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_stream = BytesIO(pix.tobytes("png"))

                self.doc.add_picture(img_stream, width=Inches(6.2))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_paragraph.paragraph_format.space_after = Pt(2)

            pdf_doc.close()

    def save(self, filename: str):
        """Salva o documento DOCX."""
        self.doc.save(filename)
