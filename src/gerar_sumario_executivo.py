# -*- coding: utf-8 -*-
"""
Gerador do Sumário Executivo — Curso ABAR de Análise de Dados de Gás
Gera documento DOCX profissional com screenshots da interface web e gráficos do projeto.
"""

import os
import sys
import time
import subprocess
import signal

from pathlib import Path

# Ensure src/ is on path
SRC_DIR = Path(__file__).resolve().parent
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from config import PROJECT_ROOT, GRAFICOS_DIR, DIAGRAMAS_DIR, OUTPUTS_DIR

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# === CONFIGURAÇÃO ===
LOGO_DIR = Path(r"C:\Users\raula\OneDrive\Documentos\Codigos e sistemas pessoais\AGENERSA\Agenersa Dashbborad Rede\Logo")
SCREENSHOT_DIR = OUTPUTS_DIR / "screenshots"
OUTPUT_DIR = Path(r"C:\Users\raula\OneDrive\Documentos\Codigos e sistemas pessoais\AGENERSA\MVPs\Sumarios")
OUTPUT_FILE = OUTPUT_DIR / "Sumario_Executivo_Curso_ABAR_Dados.docx"

# Cores institucionais
COR_PRINCIPAL = RGBColor(0x1B, 0x3A, 0x5C)   # #1B3A5C
COR_SECUNDARIA = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6
COR_CINZA = RGBColor(0x59, 0x56, 0x59)         # #595659
COR_BRANCA = RGBColor(0xFF, 0xFF, 0xFF)

IMG_WIDTH = Cm(16)
LOGO_HEIGHT = Cm(1.8)


# =====================================================================
# HELPERS (padrão DataReg360)
# =====================================================================

def configurar_estilos(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COR_CINZA
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level, size in [(1, 16), (2, 13), (3, 11)]:
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = COR_PRINCIPAL
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)


def adicionar_paragrafo(doc, texto, bold=False, italic=False, size=None, color=None,
                        alignment=None, space_before=None, space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(texto)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = 'Calibri'
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def adicionar_bullet(doc, texto, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(11)
        run_bold.font.color.rgb = COR_CINZA
        run_rest = p.add_run(texto)
        run_rest.font.name = 'Calibri'
        run_rest.font.size = Pt(11)
        run_rest.font.color.rgb = COR_CINZA
    else:
        run = p.add_run(texto)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = COR_CINZA
    return p


def adicionar_imagem(doc, caminho, legenda=None, largura=None):
    caminho = str(caminho)
    if not os.path.exists(caminho):
        print(f"  [AVISO] Imagem não encontrada: {caminho}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(caminho, width=largura or IMG_WIDTH)
    if legenda:
        p_leg = doc.add_paragraph()
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_leg = p_leg.add_run(legenda)
        run_leg.italic = True
        run_leg.font.size = Pt(9)
        run_leg.font.color.rgb = COR_CINZA
        run_leg.font.name = 'Calibri'
        p_leg.paragraph_format.space_after = Pt(12)


def adicionar_tabela(doc, cabecalhos, linhas):
    table = doc.add_table(rows=1, cols=len(cabecalhos))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, cab in enumerate(cabecalhos):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cab)
        run.bold = True
        run.font.color.rgb = COR_BRANCA
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for linha in linhas:
        row = table.add_row()
        for i, valor in enumerate(linha):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(valor)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = COR_CINZA

    doc.add_paragraph()
    return table


# =====================================================================
# CAPTURA DE SCREENSHOTS (Selenium headless)
# =====================================================================

def capturar_screenshots_web():
    """Inicia o servidor FastAPI, captura 6 screenshots via Selenium, encerra."""
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)

    # Iniciar o servidor FastAPI como subprocess
    print("  Iniciando servidor FastAPI...")
    server_proc = subprocess.Popen(
        [sys.executable, str(SRC_DIR / "run_web.py")],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        cwd=str(PROJECT_ROOT),
    )
    time.sleep(4)  # Aguardar startup

    try:
        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options

        opts = Options()
        opts.add_argument("--headless=new")
        opts.add_argument("--window-size=1600,900")
        opts.add_argument("--force-device-scale-factor=1")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--no-sandbox")

        driver = webdriver.Chrome(options=opts)
        driver.get("http://127.0.0.1:8000")
        time.sleep(3)

        screenshots = []

        # 1. Config (página inicial)
        print("  Screenshot 1/6: Configuração...")
        path1 = str(SCREENSHOT_DIR / "01_web_config.png")
        driver.save_screenshot(path1)
        screenshots.append(path1)

        # 2. Pipeline
        print("  Screenshot 2/6: Pipeline...")
        nav = driver.find_element(By.CSS_SELECTOR, '[data-section="pipeline"]')
        nav.click()
        time.sleep(1)
        path2 = str(SCREENSHOT_DIR / "02_web_pipeline.png")
        driver.save_screenshot(path2)
        screenshots.append(path2)

        # 3. Gráficos
        print("  Screenshot 3/6: Gráficos...")
        nav = driver.find_element(By.CSS_SELECTOR, '[data-section="graficos"]')
        nav.click()
        time.sleep(2)
        path3 = str(SCREENSHOT_DIR / "03_web_graficos.png")
        driver.save_screenshot(path3)
        screenshots.append(path3)

        # 4. Diagramas
        print("  Screenshot 4/6: Diagramas...")
        nav = driver.find_element(By.CSS_SELECTOR, '[data-section="diagramas"]')
        nav.click()
        time.sleep(2)
        path4 = str(SCREENSHOT_DIR / "04_web_diagramas.png")
        driver.save_screenshot(path4)
        screenshots.append(path4)

        # 5. Textos
        print("  Screenshot 5/6: Textos...")
        nav = driver.find_element(By.CSS_SELECTOR, '[data-section="textos"]')
        nav.click()
        time.sleep(2)
        path5 = str(SCREENSHOT_DIR / "05_web_textos.png")
        driver.save_screenshot(path5)
        screenshots.append(path5)

        # 6. Downloads
        print("  Screenshot 6/6: Downloads...")
        nav = driver.find_element(By.CSS_SELECTOR, '[data-section="downloads"]')
        nav.click()
        time.sleep(1)
        path6 = str(SCREENSHOT_DIR / "06_web_downloads.png")
        driver.save_screenshot(path6)
        screenshots.append(path6)

        driver.quit()
        print(f"  {len(screenshots)} screenshots salvos em {SCREENSHOT_DIR}")

    finally:
        print("  Encerrando servidor...")
        server_proc.terminate()
        try:
            server_proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            server_proc.kill()

    return screenshots


# =====================================================================
# SEÇÕES DO DOCUMENTO
# =====================================================================

def criar_capa(doc):
    for _ in range(2):
        adicionar_paragrafo(doc, '', space_after=0)

    # Logos
    logos = [
        ("Logo ABAR.jpeg", LOGO_HEIGHT),
        ("Logo Lab DGE.jpg", LOGO_HEIGHT),
        ("Logo UFF.jpeg", LOGO_HEIGHT),
    ]
    p_logos = doc.add_paragraph()
    p_logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for nome, altura in logos:
        caminho = LOGO_DIR / nome
        if caminho.exists():
            run = p_logos.add_run()
            run.add_picture(str(caminho), height=altura)
            run.add_text("   ")

    adicionar_paragrafo(doc, '', space_after=24)

    # Linha separadora
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run('_' * 60)
    run_line.font.color.rgb = COR_SECUNDARIA
    run_line.font.size = Pt(8)

    adicionar_paragrafo(doc, '', space_after=12)

    # Título
    adicionar_paragrafo(
        doc,
        'Análise Automatizada de Dados\nde Distrito de Gás',
        bold=True, size=28, color=COR_PRINCIPAL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    adicionar_paragrafo(
        doc,
        'Curso ABAR — Medições Inteligentes\ne Gestão Integrada',
        bold=True, size=18, color=COR_SECUNDARIA,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )

    # Linha separadora
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run('_' * 60)
    run_line2.font.color.rgb = COR_SECUNDARIA
    run_line2.font.size = Pt(8)

    adicionar_paragrafo(doc, '', space_after=12)

    adicionar_paragrafo(
        doc,
        'Sumário Executivo',
        bold=True, size=16, color=COR_PRINCIPAL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    adicionar_paragrafo(
        doc,
        'Uso de Inteligência Artificial para Geração de\nRelatórios de Auditoria e Análise de Dados',
        italic=True, size=13, color=COR_SECUNDARIA,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30
    )

    # Créditos
    creditos = [
        ('Coordenação:', ' Vladimir Paschoal Macedo'),
        ('Orientação:', ' Prof. Alexandre Beraldi Santos'),
        ('Autor:', ' Raul Araújo da Silva'),
    ]
    for label, valor in creditos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = COR_PRINCIPAL
        run_label.font.name = 'Calibri'
        run_valor = p.add_run(valor)
        run_valor.font.size = Pt(11)
        run_valor.font.color.rgb = COR_CINZA
        run_valor.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(2)

    adicionar_paragrafo(doc, '', space_after=18)

    adicionar_paragrafo(
        doc,
        'Fevereiro de 2026',
        bold=True, size=12, color=COR_PRINCIPAL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    doc.add_page_break()


def secao_introducao(doc):
    doc.add_heading('1. Introdução', level=1)

    adicionar_paragrafo(
        doc,
        'No âmbito do curso "Medições Inteligentes e Gestão Integrada" promovido pela '
        'ABAR (Associação Brasileira de Agências de Regulação), a AGENERSA — Agência '
        'Reguladora de Energia e Saneamento Básico do Estado do Rio de Janeiro — '
        'desenvolveu um projeto piloto de análise automatizada de dados de distrito de gás, '
        'utilizando Inteligência Artificial em todas as etapas do processo.'
    )

    adicionar_paragrafo(
        doc,
        'O objetivo do projeto é demonstrar como a IA pode transformar uma planilha Excel '
        'bruta com dados operacionais de um distrito de gás em um relatório técnico de '
        'auditoria completo, com gráficos, tabelas, equações e pareceres regulatórios — '
        'de forma automatizada, reproduzível e auditável.'
    )

    adicionar_paragrafo(
        doc,
        'As entregas do projeto incluem: 7 notebooks Jupyter de análise, 23 gráficos '
        'com qualidade de publicação, 3 diagramas de processo gerados por IA, relatório '
        'DOCX de auditoria com 9.3 MB (incluindo equações nativas e apêndice com código), '
        'apresentação PPTX auto-gerada com 13 slides, e interface web para monitoramento '
        'em tempo real do pipeline de geração. Os dados analisados compreendem 183 dias de '
        'operação de 1 distrito com 7 clientes industriais.'
    )


def secao_problema(doc):
    doc.add_heading('2. O Problema: Limitações da Análise Manual', level=1)

    adicionar_paragrafo(
        doc,
        'A análise de dados operacionais de distritos de gás enfrenta desafios '
        'significativos quando realizada de forma manual:'
    )

    problemas = [
        ('Complexidade dos dados: ',
         'planilhas Excel com 14 abas contendo volumes, pressões, temperaturas, PCS '
         'e dados de 7 clientes industriais.'),
        ('Conhecimento especializado: ',
         'a análise requer domínio de metodologias como GUM (Guide to the Expression of '
         'Uncertainty in Measurement), propagação RSS, cálculo de PCS e balanço de massa.'),
        ('Falta de reprodutibilidade: ',
         'análises feitas manualmente em Excel não são auditáveis nem facilmente '
         'reproduzíveis por terceiros.'),
        ('Ausência de padronização: ',
         'não existe estrutura padronizada para relatórios técnicos de auditoria de '
         'dados de distrito.'),
        ('Esforço de integração: ',
         'cruzar teoria (metodologia) com dados numéricos e produzir texto técnico '
         'coerente é trabalhoso e demanda tempo significativo.'),
    ]
    for bold_part, rest in problemas:
        adicionar_bullet(doc, rest, bold_prefix=bold_part)


def secao_solucao(doc):
    doc.add_heading('3. A Solução: Pipeline Automatizado com IA', level=1)

    adicionar_paragrafo(
        doc,
        'O projeto implementa um pipeline completo em Python que combina análise de dados '
        'via notebooks, extração de metodologia via IA e geração de relatório detalhado '
        'via LLM (Large Language Model), produzindo múltiplos formatos de saída.'
    )

    # 3.1
    doc.add_heading('3.1 Extração de Metodologia via IA', level=2)

    adicionar_paragrafo(
        doc,
        'A IA (Google Gemini) recebe a apostila do curso em PDF e extrai automaticamente '
        'a metodologia técnica, produzindo 7 textos em Markdown e um arquivo JSON com '
        'equações. A IA não inventa conteúdo — ela estrutura e organiza o conhecimento '
        'existente na apostila. Além disso, gera 3 diagramas de processo via Gemini Image '
        'Preview, ilustrando o fluxo da auditoria, a estrutura do distrito e o processo '
        'de análise.'
    )

    # 3.2
    doc.add_heading('3.2 Análise de Dados Reproduzível via Notebooks', level=2)

    adicionar_paragrafo(
        doc,
        'O projeto inclui 7 Jupyter Notebooks que realizam a análise completa dos dados '
        'do Excel, desde a leitura e exploração até o balanço de massa final. O código é '
        'visível, auditável e reproduzível. Os notebooks geram 23 gráficos com qualidade '
        'de publicação usando matplotlib, cobrindo: volumes de entrada, PCS (Poder '
        'Calorífico Superior), energia, perfis de consumo dos clientes, incertezas de '
        'medição (metodologia GUM) e balanço de massa com bandas de incerteza.'
    )

    # 3.3
    doc.add_heading('3.3 Geração de Relatório Detalhado via LLM', level=2)

    adicionar_paragrafo(
        doc,
        'O pipeline de geração do relatório realiza 28 chamadas estruturadas ao LLM '
        '(4 por capítulo): Metodologia, Dados Numéricos, Gráficos e Síntese. O modelo '
        'recebe como contexto os dados extraídos dos notebooks, a metodologia da apostila '
        'e os gráficos gerados, produzindo texto técnico detalhado para cada capítulo.'
    )

    adicionar_paragrafo(
        doc,
        'Cada capítulo segue a estrutura: Introdução → Fundamentação Teórica → Análise '
        'de Dados → Discussão dos Gráficos → Parecer Regulatório. O sistema mantém um '
        'cache de 28 arquivos Markdown que permite remontar o relatório sem novas chamadas '
        'à API. A montagem final produz um DOCX com equações nativas (OMML), 6 tabelas '
        'estatísticas, 26 imagens inline e um apêndice com o código dos 7 notebooks.'
    )

    # Diagrama de fluxo em texto
    adicionar_paragrafo(doc, '', space_after=6)
    p_flow = doc.add_paragraph()
    p_flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_flow = p_flow.add_run(
        'Excel → Notebooks → 23 PNGs + dados → Gemini (28 chamadas) → '
        'Cache .md → DOCX Builder → Relatório 9.3 MB'
    )
    run_flow.bold = True
    run_flow.font.size = Pt(10)
    run_flow.font.color.rgb = COR_SECUNDARIA
    run_flow.font.name = 'Calibri'
    p_flow.paragraph_format.space_after = Pt(12)


def secao_interface_web(doc):
    doc.add_heading('4. Interface Web de Monitoramento', level=1)

    adicionar_paragrafo(
        doc,
        'Para facilitar a execução e o acompanhamento do pipeline, foi desenvolvida uma '
        'interface web local utilizando FastAPI no backend e JavaScript vanilla no frontend, '
        'com comunicação em tempo real via Server-Sent Events (SSE). A interface permite '
        'upload do arquivo Excel, configuração da API key, seleção do modo de execução, '
        'e acompanha as 29 etapas do pipeline com barra de progresso, stepper de fases '
        'e cards por capítulo.'
    )

    # 4.1 Config
    doc.add_heading('4.1 Configuração e Upload', level=2)

    adicionar_paragrafo(
        doc,
        'A tela inicial permite o upload do arquivo Excel com os dados do distrito, '
        'a configuração da chave da API Gemini e a seleção do modo de execução '
        '(geração completa ou apenas montagem a partir do cache).'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "01_web_config.png",
        'Figura 1 — Interface web: tela de configuração e upload'
    )

    # 4.2 Pipeline
    doc.add_heading('4.2 Pipeline em Tempo Real', level=2)

    adicionar_paragrafo(
        doc,
        'Durante a execução, a interface exibe uma barra de progresso global, um stepper '
        'com as 4 fases do pipeline e cards detalhados para cada capítulo, mostrando '
        'o status de cada uma das 29 etapas em tempo real.'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "02_web_pipeline.png",
        'Figura 2 — Interface web: monitoramento do pipeline em tempo real'
    )

    # 4.3 Gráficos
    doc.add_heading('4.3 Galeria de Gráficos', level=2)

    adicionar_paragrafo(
        doc,
        'Após a execução, os 23 gráficos gerados pelos notebooks são exibidos em uma '
        'galeria organizada por capítulo, com tabs de navegação.'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "03_web_graficos.png",
        'Figura 3 — Interface web: galeria de gráficos por capítulo'
    )

    # 4.4 Diagramas
    doc.add_heading('4.4 Diagramas de Processo', level=2)

    adicionar_paragrafo(
        doc,
        'Os 3 diagramas gerados pela IA (Gemini Image) são exibidos em seção dedicada: '
        'fluxo da auditoria, estrutura do distrito e processo de análise.'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "04_web_diagramas.png",
        'Figura 4 — Interface web: diagramas de processo gerados por IA'
    )

    # 4.5 Textos
    doc.add_heading('4.5 Textos Gerados', level=2)

    adicionar_paragrafo(
        doc,
        'Os textos gerados pelo LLM para cada capítulo são exibidos em formato acordeão '
        'com renderização Markdown, permitindo revisão rápida do conteúdo.'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "05_web_textos.png",
        'Figura 5 — Interface web: textos gerados pelo LLM em acordeão'
    )

    # 4.6 Downloads
    doc.add_heading('4.6 Downloads', level=2)

    adicionar_paragrafo(
        doc,
        'A seção de downloads oferece acesso direto aos arquivos gerados: relatório DOCX '
        'de auditoria e apresentação PPTX, com informações de tamanho de cada arquivo.'
    )

    adicionar_imagem(
        doc,
        SCREENSHOT_DIR / "06_web_downloads.png",
        'Figura 6 — Interface web: download do relatório e apresentação'
    )


def secao_exemplos(doc):
    doc.add_heading('5. Exemplos de Saída', level=1)

    # 5.1 Gráficos
    doc.add_heading('5.1 Gráficos de Análise', level=2)

    adicionar_paragrafo(
        doc,
        'A seguir, quatro gráficos representativos das análises realizadas pelos '
        'notebooks Jupyter, demonstrando a qualidade e abrangência das visualizações:'
    )

    graficos = [
        ("vol_entrada_serie.png", "Figura 7 — Série temporal dos volumes de entrada (Nm³/d)"),
        ("clientes_participacao.png", "Figura 8 — Participação dos 7 clientes no consumo total"),
        ("balanco_waterfall.png", "Figura 9 — Balanço de massa com bandas de incerteza (waterfall)"),
        ("incertezas_barras.png", "Figura 10 — Incertezas combinadas por ponto de medição"),
    ]
    for arquivo, legenda in graficos:
        adicionar_imagem(doc, GRAFICOS_DIR / arquivo, legenda)

    # 5.2 Diagramas
    doc.add_heading('5.2 Diagramas Gerados pela IA', level=2)

    adicionar_paragrafo(
        doc,
        'Os diagramas abaixo foram gerados automaticamente pelo modelo Gemini Image '
        'Preview a partir de descrições textuais, ilustrando processos técnicos do '
        'distrito de gás:'
    )

    diagramas = [
        ("fluxo_auditoria.png", "Figura 11 — Fluxo da auditoria técnica (gerado pelo Gemini)"),
        ("estrutura_distrito.png", "Figura 12 — Estrutura do distrito analisado"),
    ]
    for arquivo, legenda in diagramas:
        adicionar_imagem(doc, DIAGRAMAS_DIR / arquivo, legenda)


def secao_arquitetura(doc):
    doc.add_heading('6. Arquitetura Técnica', level=1)

    adicionar_paragrafo(
        doc,
        'O sistema foi desenvolvido inteiramente em Python, com uma arquitetura modular '
        'que separa análise de dados, geração de conteúdo e construção de documentos:'
    )

    adicionar_tabela(doc, ['Componente', 'Tecnologia'], [
        ['Notebooks de Análise', 'Python, pandas, NumPy, matplotlib'],
        ['Modelo de IA (Texto)', 'Google Gemini 3 Pro Preview'],
        ['Modelo de IA (Imagem)', 'Google Gemini 3 Pro Image Preview'],
        ['Construtor de Relatório', 'python-docx, latex2mathml, lxml'],
        ['Construtor de Apresentação', 'python-pptx'],
        ['Interface Web', 'FastAPI, Jinja2, JavaScript (SPA)'],
        ['Streaming de Progresso', 'Server-Sent Events (SSE)'],
        ['Dados de Entrada', 'Excel / openpyxl'],
    ])

    adicionar_paragrafo(
        doc,
        'O fluxo completo segue seis fases: (1) extração de metodologia da apostila via IA; '
        '(2) geração de diagramas de processo via Gemini Image; (3) análise de dados em '
        '7 notebooks Jupyter; (4) geração segmentada de texto via 28 chamadas ao LLM; '
        '(5) montagem do DOCX com equações OMML, tabelas e gráficos; (6) geração da '
        'apresentação PPTX com 13 slides.'
    )


def secao_papel_ia(doc):
    doc.add_heading('7. O Papel da IA na Análise de Dados', level=1)

    adicionar_paragrafo(
        doc,
        'A Inteligência Artificial desempenhou um papel central em todas as etapas deste '
        'projeto, demonstrando como a tecnologia pode transformar a atuação de agências '
        'reguladoras na análise de dados técnicos. A seguir, destacam-se as cinco '
        'dimensões de atuação da IA neste projeto:'
    )

    # 7.1
    doc.add_heading('7.1 IA como Extratora de Metodologia', level=2)

    adicionar_paragrafo(
        doc,
        'A IA foi utilizada para ler a apostila do curso (PDF) e extrair automaticamente '
        'toda a fundamentação teórica necessária para a auditoria: conceitos de PCS, '
        'metodologia GUM para incertezas, propagação RSS, balanço de massa e critérios '
        'regulatórios. O resultado foram 7 textos estruturados em Markdown e um arquivo '
        'JSON com equações — tudo organizado de forma a servir como referência para o '
        'pipeline de geração do relatório. A IA não inventa conteúdo; ela estrutura e '
        'organiza o conhecimento existente.'
    )

    # 7.2
    doc.add_heading('7.2 IA como Geradora de Código', level=2)

    adicionar_paragrafo(
        doc,
        'Todo o código Python do projeto — incluindo os 7 notebooks de análise, os '
        'scripts de geração de relatório e apresentação, e a interface web — foi '
        'desenvolvido com assistência de IA (Claude Code, Anthropic). A IA auxiliou na '
        'leitura e processamento dos dados Excel, na implementação das análises '
        'estatísticas, na geração de gráficos com matplotlib e na construção da '
        'aplicação web com FastAPI. Isso permitiu que um único analista produzisse '
        'um sistema completo que normalmente demandaria uma equipe de desenvolvimento.'
    )

    # 7.3
    doc.add_heading('7.3 IA como Analista Detalhada', level=2)

    adicionar_paragrafo(
        doc,
        'Na geração do relatório, o LLM recebe três tipos de contexto para cada capítulo: '
        'a metodologia extraída da apostila, os dados numéricos dos notebooks e os gráficos '
        'gerados. Com esse contexto estruturado, a IA produz texto técnico detalhado — '
        'incluindo fundamentação teórica, análise dos dados, discussão dos gráficos e '
        'pareceres regulatórios. Cada capítulo é gerado em 4 chamadas especializadas '
        '(Metodologia, Dados, Gráficos, Síntese), totalizando 28 chamadas que produzem '
        'um relatório abrangente e coerente.'
    )

    # 7.4
    doc.add_heading('7.4 IA como Produtora de Múltiplos Formatos', level=2)

    adicionar_paragrafo(
        doc,
        'O pipeline demonstra a capacidade da IA de produzir saídas em múltiplos formatos '
        'profissionais: relatório DOCX com equações nativas (OMML), tabelas formatadas e '
        '26 imagens inline; apresentação PPTX auto-gerada com 13 slides e gráficos; '
        'diagramas PNG gerados via Gemini Image Preview; e 23 gráficos de análise via '
        'matplotlib. Todos os formatos são gerados automaticamente a partir dos mesmos '
        'dados e contexto, garantindo consistência.'
    )

    # 7.5
    doc.add_heading('7.5 IA como Multiplicadora de Capacidade Regulatória', level=2)

    adicionar_paragrafo(
        doc,
        'Para agências reguladoras com equipes técnicas limitadas, a IA representa um '
        'multiplicador de capacidade sem precedentes. O pipeline desenvolvido neste projeto '
        'transforma o que seria semanas de trabalho manual — leitura de dados, análise '
        'estatística, cruzamento com metodologia, redação de relatório técnico — em um '
        'processo automatizado que executa em aproximadamente 18 minutos (geração completa) '
        'ou 7 segundos (remontagem a partir do cache). Todo o processo é transparente: o '
        'código dos notebooks é visível e auditável, os textos gerados ficam em cache .md, '
        'e o pipeline pode ser re-executado por qualquer analista.'
    )


def secao_resultados(doc):
    doc.add_heading('8. Resultados Obtidos', level=1)

    adicionar_paragrafo(
        doc,
        'O projeto alcançou os seguintes resultados concretos:'
    )

    resultados = [
        '183 dias de dados processados em 7 notebooks analíticos',
        '23 gráficos + 3 diagramas IA gerados com qualidade de publicação',
        'Relatório de auditoria DOCX (9.3 MB) com equações nativas, 6 tabelas, 26 imagens e apêndice com código',
        'Apresentação PPTX auto-gerada com 13 slides e gráficos',
        'Interface web com monitoramento em tempo real de 29 etapas do pipeline',
        'Cache com 28 arquivos permite reconstrução completa sem chamadas à API',
        'Anonimização de 7 empresas (Empresa A–G) para proteção de dados',
        'Custo de API: aproximadamente zero (tier gratuito do Google Gemini)',
        'Tempo do pipeline: ~18 min (geração completa) ou ~7s (apenas montagem)',
    ]
    for r in resultados:
        adicionar_bullet(doc, r)


def secao_proximos_passos(doc):
    doc.add_heading('9. Próximos Passos', level=1)

    adicionar_paragrafo(
        doc,
        'O pipeline desenvolvido pode ser expandido e aplicado em cenários reais:'
    )

    passos = [
        ('Aplicação a dados reais: ',
         'utilizar dados reais de distritos das concessões da AGENERSA.'),
        ('Integração periódica: ',
         'conectar o pipeline a fluxos periódicos de coleta de dados operacionais.'),
        ('Expansão temática: ',
         'incluir capítulos adicionais sobre tarifas, qualidade de serviço e '
         'indicadores de desempenho.'),
        ('Deploy da interface: ',
         'disponibilizar a interface web para uso interno da agência.'),
        ('Outros setores: ',
         'expandir a abordagem para água, saneamento e outros serviços regulados.'),
    ]
    for bold_part, rest in passos:
        adicionar_bullet(doc, rest, bold_prefix=bold_part)


def secao_conclusao(doc):
    doc.add_heading('10. Conclusão', level=1)

    adicionar_paragrafo(
        doc,
        'Este projeto demonstra que a Inteligência Artificial pode transformar '
        'completamente o processo de auditoria de dados de distritos de gás. O pipeline '
        'desenvolvido transforma um arquivo Excel bruto em um relatório profissional de '
        'auditoria — com gráficos, equações, tabelas e pareceres regulatórios — de forma '
        'automatizada e reproduzível.'
    )

    adicionar_paragrafo(
        doc,
        'A combinação de notebooks Jupyter (análise auditável), extração de metodologia '
        'via IA (fundamentação teórica) e geração de texto via LLM (relatório detalhado) '
        'cria um fluxo de trabalho escalável: o mesmo pipeline pode ser aplicado a '
        'qualquer distrito de gás, bastando fornecer o arquivo Excel correspondente.'
    )

    adicionar_paragrafo(
        doc,
        'O impacto para agências reguladoras é direto: um analista com acesso ao pipeline '
        'pode produzir relatórios técnicos que antes exigiriam semanas de trabalho, '
        'com qualidade consistente e total rastreabilidade. A IA não substitui o '
        'julgamento humano — ela amplifica a capacidade técnica disponível.'
    )


def secao_creditos(doc):
    doc.add_page_break()
    doc.add_heading('Créditos', level=1)

    creditos = [
        ('Coordenador:', 'Vladimir Paschoal Macedo'),
        ('Orientador:', 'Prof. Alexandre Beraldi Santos'),
        ('Autor:', 'Raul Araújo da Silva'),
        ('Parceria Institucional:', 'ABAR – Associação Brasileira de Agências de Regulação'),
        ('Vinculação Acadêmica:', 'LabDGE – Laboratório de Dados e Governança Energética / UFF'),
    ]

    for label, valor in creditos:
        p = doc.add_paragraph()
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = COR_PRINCIPAL
        run_label.font.name = 'Calibri'
        run_valor = p.add_run(valor)
        run_valor.font.size = Pt(11)
        run_valor.font.color.rgb = COR_CINZA
        run_valor.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(4)

    adicionar_paragrafo(doc, '', space_after=24)

    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nota = p_nota.add_run(
        'Este documento e o projeto foram desenvolvidos com assistência de '
        'Inteligência Artificial (Claude, Anthropic e Gemini, Google).'
    )
    run_nota.italic = True
    run_nota.font.size = Pt(9)
    run_nota.font.color.rgb = COR_CINZA
    run_nota.font.name = 'Calibri'


# =====================================================================
# MAIN
# =====================================================================

def main():
    print("=" * 60)
    print("  Gerador de Sumário Executivo — Curso ABAR de Dados")
    print("=" * 60)

    # Passo 1: Capturar screenshots
    print("\n[1/2] Capturando screenshots da interface web...")
    try:
        capturar_screenshots_web()
    except Exception as e:
        print(f"  [AVISO] Erro ao capturar screenshots: {e}")
        print("  Continuando sem screenshots da web...")

    # Passo 2: Gerar documento
    print("\n[2/2] Gerando documento DOCX...")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    configurar_estilos(doc)

    print("  [1/12] Capa...")
    criar_capa(doc)

    print("  [2/12] Seção 1: Introdução...")
    secao_introducao(doc)

    print("  [3/12] Seção 2: O Problema...")
    secao_problema(doc)

    print("  [4/12] Seção 3: A Solução...")
    secao_solucao(doc)

    print("  [5/12] Seção 4: Interface Web...")
    secao_interface_web(doc)

    print("  [6/12] Seção 5: Exemplos de Saída...")
    secao_exemplos(doc)

    print("  [7/12] Seção 6: Arquitetura Técnica...")
    secao_arquitetura(doc)

    print("  [8/12] Seção 7: Papel da IA...")
    secao_papel_ia(doc)

    print("  [9/12] Seção 8: Resultados...")
    secao_resultados(doc)

    print("  [10/12] Seção 9: Próximos Passos...")
    secao_proximos_passos(doc)

    print("  [11/12] Seção 10: Conclusão...")
    secao_conclusao(doc)

    print("  [12/12] Créditos...")
    secao_creditos(doc)

    doc.save(str(OUTPUT_FILE))

    print(f"\n{'=' * 60}")
    print(f"  Documento gerado com sucesso!")
    print(f"  Caminho: {OUTPUT_FILE}")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()
