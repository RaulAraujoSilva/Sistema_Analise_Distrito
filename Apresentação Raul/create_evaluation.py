# -*- coding: utf-8 -*-
"""
Script para criar documentos Word de avaliação — Bloco 6 ABAR
Curso: Infraestrutura de Gás — Medições Inteligentes e Gestão Integrada
"""

import os
import shutil
import zipfile
from copy import deepcopy
from lxml import etree

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Caminhos ──────────────────────────────────────────────────────────────────
BASE_DIR = r"C:\Users\raula\OneDrive\Documentos\Codigos e sistemas pessoais\AGENERSA\Cursos ABAR de Dados\Apresentação Raul"
TEMPLATE_PATH = os.path.join(BASE_DIR, "folha timbrada ABAR.docx")
OUT_ALUNO    = os.path.join(BASE_DIR, "Avaliacao_Bloco6_aluno.docx")
OUT_PROF     = os.path.join(BASE_DIR, "Avaliacao_Bloco6_professor.docx")

# ── Questões e gabarito ───────────────────────────────────────────────────────
GABARITO = {1: 'b', 2: 'c', 3: 'b', 4: 'c', 5: 'b',
            6: 'c', 7: 'b', 8: 'b', 9: 'b', 10: 'b'}

QUESTOES = [
    {
        "num": 1,
        "enunciado": "Qual das afirmações abaixo melhor descreve uma conquista recente de modelos de IA em 2025?",
        "alternativas": [
            ("a", "Modelos de IA atingiram pela primeira vez pontuação acima de 50% no benchmark MMLU"),
            ("b", "Um modelo de IA obteve medalha de ouro na Olimpíada Internacional de Matemática (IMO)"),
            ("c", "O primeiro chatbot comercial foi lançado ao público"),
            ("d", "IA superou humanos apenas em jogos de tabuleiro como xadrez"),
        ],
    },
    {
        "num": 2,
        "enunciado": "Um regulador sem conhecimento de programação precisa analisar uma planilha simples de dados mensais. Qual ferramenta é mais indicada?",
        "alternativas": [
            ("a", "Claude Code CLI"),
            ("b", "VS Code com GitHub Copilot"),
            ("c", "ChatGPT na web"),
            ("d", "TensorFlow Playground"),
        ],
    },
    {
        "num": 3,
        "enunciado": "Na IBM Data Science Methodology, qual etapa é potencializada quando usamos IA para ler documentos técnicos (como apostilas e normas) e extrair automaticamente equações e metodologias?",
        "alternativas": [
            ("a", "Coleta dos Dados"),
            ("b", "Entendimento do Negócio / Criação do Modelo"),
            ("c", "Deployment"),
            ("d", "Feedback"),
        ],
    },
    {
        "num": 4,
        "enunciado": "Ao pedir ao Claude Code para coletar dados de uma agência reguladora na internet, estamos aplicando qual etapa da metodologia de dados?",
        "alternativas": [
            ("a", "Criação do Modelo"),
            ("b", "Avaliação"),
            ("c", "Coleta dos Dados"),
            ("d", "Preparação dos Dados"),
        ],
    },
    {
        "num": 5,
        "enunciado": "O sistema demonstrado nesta aula parte de uma planilha Excel e gera um relatório Word de mais de 100 páginas em menos de 2 minutos. Qual tecnologia é responsável pelo conteúdo textual do relatório?",
        "alternativas": [
            ("a", "Template Word preenchido manualmente"),
            ("b", "Modelo de linguagem (LLM) Gemini chamado via API"),
            ("c", "Algoritmo de machine learning treinado com dados históricos"),
            ("d", "Macro VBA do Excel"),
        ],
    },
    {
        "num": 6,
        "enunciado": "Para que serve o NotebookLM no contexto desta aula?",
        "alternativas": [
            ("a", "Executar notebooks Jupyter na nuvem"),
            ("b", "Criar dashboards interativos a partir de dados"),
            ("c", "Carregar documentos técnicos e conversar com eles usando IA"),
            ("d", "Substituir o Python na análise de dados"),
        ],
    },
    {
        "num": 7,
        "enunciado": "O sistema de auditoria demonstrado reduziu o tempo de geração do relatório de ~20 minutos para ~1 minuto e 42 segundos. Qual estratégia foi usada?",
        "alternativas": [
            ("a", "Uso de hardware mais potente (GPU)"),
            ("b", "Chamadas paralelas à API de IA (28 seções simultâneas)"),
            ("c", "Caching de dados no navegador"),
            ("d", "Compressão do arquivo Word gerado"),
        ],
    },
    {
        "num": 8,
        "enunciado": "Um auditor regulatório usa IA para verificar se o balanço de massa de um distrito de gás está equilibrado, usando dados de 183 dias. Isso se enquadra em qual categoria de aplicação?",
        "alternativas": [
            ("a", "IA gerativa para criação de conteúdo artístico"),
            ("b", "IA aplicada à transparência regulatória e conformidade técnica"),
            ("c", "Machine learning preditivo para previsão de demanda"),
            ("d", "Processamento de linguagem natural para análise de sentimentos"),
        ],
    },
    {
        "num": 9,
        "enunciado": "Qual é o principal argumento desta aula em favor do uso de IA na ciência de dados para reguladores e auditores?",
        "alternativas": [
            ("a", "IA substitui completamente a necessidade de conhecimento técnico no domínio"),
            ("b", "IA permite que profissionais com conhecimento de domínio (gás, regulação) executem análises sem precisar programar"),
            ("c", "IA é mais precisa do que métodos estatísticos tradicionais em qualquer contexto"),
            ("d", "Ferramentas de IA são gratuitas e não exigem infraestrutura"),
        ],
    },
    {
        "num": 10,
        "enunciado": "No projeto demonstrado nesta aula, qual é a ordem correta do fluxo de trabalho?",
        "alternativas": [
            ("a", "Relatório Word → Planilha Excel → Notebooks → API Gemini"),
            ("b", "Planilha Excel → 7 Notebooks de análise → API Gemini (28 chamadas) → Relatório Word (100+ pág.)"),
            ("c", "Apostila PDF → ChatGPT → Planilha Excel → Dashboard"),
            ("d", "Dashboard → Notebooks → Planilha Excel → Relatório"),
        ],
    },
]

# ── Cores ABAR ────────────────────────────────────────────────────────────────
ABAR_AZUL      = RGBColor(0x00, 0x38, 0x86)   # azul institucional
ABAR_VERDE     = RGBColor(0x00, 0x80, 0x00)   # verde para gabarito
COR_GABARITO   = RGBColor(0x1F, 0x5C, 0xAB)  # azul destaque para gabarito professor
CINZA_CLARO    = RGBColor(0xF2, 0xF2, 0xF2)


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_paragraph_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    # Remove existing spacing element if any
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing)


def add_horizontal_line(doc, color_hex="003886", width_pct=100):
    """Adiciona uma linha horizontal colorida."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=0, after=0)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def shade_paragraph(para, fill_hex):
    """Aplica cor de fundo em um parágrafo."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    for existing in pPr.findall(qn('w:shd')):
        pPr.remove(existing)
    pPr.append(shd)


def copy_header_footer_from_template(new_doc, template_path):
    """
    Copia o cabeçalho (com imagem) e rodapé do template ABAR para o novo documento.
    Estratégia: manipulação direta dos arquivos XML dentro do ZIP.
    """
    # Salva o novo documento temporariamente
    tmp_path = template_path.replace("folha timbrada ABAR.docx", "_tmp_new.docx")
    new_doc.save(tmp_path)

    # Abre os dois ZIPs
    with zipfile.ZipFile(template_path, 'r') as tmpl_zip:
        tmpl_names = tmpl_zip.namelist()

        # Lê os arquivos de header/footer do template
        header_xml  = tmpl_zip.read('word/header1.xml')
        footer_xml  = tmpl_zip.read('word/footer1.xml')
        h_rels_xml  = tmpl_zip.read('word/_rels/header1.xml.rels')
        f_rels_xml  = tmpl_zip.read('word/_rels/footer1.xml.rels')
        image_data  = tmpl_zip.read('word/media/image1.png')
        doc_rels_xml = tmpl_zip.read('word/_rels/document.xml.rels')
        doc_xml     = tmpl_zip.read('word/document.xml')

        # Lê todos os outros arquivos do template (para reconstruir o ZIP)
        tmpl_files = {}
        for name in tmpl_names:
            tmpl_files[name] = tmpl_zip.read(name)

    with zipfile.ZipFile(tmp_path, 'r') as new_zip:
        new_names = new_zip.namelist()
        new_files = {}
        for name in new_names:
            new_files[name] = new_zip.read(name)

    # ── Modificar o document.xml do NOVO documento para referenciar header/footer ──
    new_doc_xml = new_files['word/document.xml'].decode('utf-8')
    new_doc_root = etree.fromstring(new_doc_xml.encode('utf-8'))

    # Namespaces
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    # Encontra o elemento sectPr
    body = new_doc_root.find(f'{{{W}}}body')
    sectPr = body.find(f'{{{W}}}sectPr')
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)

    # Remove referências antigas de header/footer
    for tag in [f'{{{W}}}headerReference', f'{{{W}}}footerReference']:
        for el in sectPr.findall(tag):
            sectPr.remove(el)

    # Adiciona referências ao header e footer
    hRef = OxmlElement('w:headerReference')
    hRef.set(qn('w:type'), 'default')
    hRef.set(qn('r:id'), 'rIdHdr1')
    sectPr.insert(0, hRef)

    fRef = OxmlElement('w:footerReference')
    fRef.set(qn('w:type'), 'default')
    fRef.set(qn('r:id'), 'rIdFtr1')
    sectPr.insert(1, fRef)

    new_doc_xml_modified = etree.tostring(new_doc_root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)

    # ── Modificar document.xml.rels para incluir header e footer ──
    new_doc_rels_xml = new_files['word/_rels/document.xml.rels'].decode('utf-8')
    rels_root = etree.fromstring(new_doc_rels_xml.encode('utf-8'))
    RELS_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'

    # Remove rels antigos de header/footer se existirem
    for rel in rels_root.findall(f'{{{RELS_NS}}}Relationship'):
        ttype = rel.get('Type', '')
        if 'header' in ttype or 'footer' in ttype:
            rels_root.remove(rel)

    # Adiciona novos rels
    hdr_rel = etree.SubElement(rels_root, f'{{{RELS_NS}}}Relationship')
    hdr_rel.set('Id', 'rIdHdr1')
    hdr_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
    hdr_rel.set('Target', 'header1.xml')

    ftr_rel = etree.SubElement(rels_root, f'{{{RELS_NS}}}Relationship')
    ftr_rel.set('Id', 'rIdFtr1')
    ftr_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')
    ftr_rel.set('Target', 'footer1.xml')

    new_doc_rels_modified = etree.tostring(rels_root, xml_declaration=True,
                                           encoding='UTF-8', standalone=True)

    # ── Prepara [Content_Types].xml modificado ──
    content_types_xml = new_files.get('[Content_Types].xml', b'')
    ct_root = etree.fromstring(content_types_xml)
    CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
    # Verifica se já existe Default para png
    png_exists = any(
        el.get('Extension') == 'png'
        for el in ct_root.findall(f'{{{CT_NS}}}Default')
    )
    if not png_exists:
        default_png = etree.SubElement(ct_root, f'{{{CT_NS}}}Default')
        default_png.set('Extension', 'png')
        default_png.set('ContentType', 'image/png')
    # Garante Override para header e footer
    overrides_existentes = {el.get('PartName') for el in ct_root.findall(f'{{{CT_NS}}}Override')}
    for part, ct in [
        ('/word/header1.xml', 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml'),
        ('/word/footer1.xml', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'),
    ]:
        if part not in overrides_existentes:
            ov = etree.SubElement(ct_root, f'{{{CT_NS}}}Override')
            ov.set('PartName', part)
            ov.set('ContentType', ct)
    content_types_modified = etree.tostring(ct_root, xml_declaration=True,
                                            encoding='UTF-8', standalone=True)

    # ── Reescreve o ZIP do novo documento ──
    output_path = tmp_path.replace('_tmp_new.docx', '_final_tmp.docx')
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as out_zip:
        for name, data in new_files.items():
            if name == 'word/document.xml':
                out_zip.writestr(name, new_doc_xml_modified)
            elif name == 'word/_rels/document.xml.rels':
                out_zip.writestr(name, new_doc_rels_modified)
            elif name == '[Content_Types].xml':
                out_zip.writestr(name, content_types_modified)
            else:
                out_zip.writestr(name, data)

        # Adiciona header, footer e imagem do template
        out_zip.writestr('word/header1.xml', header_xml)
        out_zip.writestr('word/footer1.xml', footer_xml)
        out_zip.writestr('word/_rels/header1.xml.rels', h_rels_xml)
        out_zip.writestr('word/_rels/footer1.xml.rels', f_rels_xml)
        # Adiciona a imagem (logo ABAR) se ainda não estiver
        if 'word/media/image1.png' not in new_files:
            out_zip.writestr('word/media/image1.png', image_data)

    # Remove temporário
    os.remove(tmp_path)

    return output_path


def build_document(versao_professor=False):
    """Constrói o documento Word de avaliação."""
    doc = Document()

    # ── Configuração da página ────────────────────────────────────────────────
    section = doc.sections[0]
    # Mesmas margens do template
    section.top_margin    = Emu(1620520)
    section.bottom_margin = Emu(900430)
    section.left_margin   = Emu(990600)
    section.right_margin  = Emu(808990)
    section.header_distance = Emu(450215)
    section.footer_distance = Emu(137795)

    # ── Estilos base ──────────────────────────────────────────────────────────
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # ── Título principal ──────────────────────────────────────────────────────
    titulo_para = doc.add_paragraph()
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(titulo_para, before=120, after=60)

    run_titulo = titulo_para.add_run("AVALIAÇÃO — BLOCO 6")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = ABAR_AZUL
    run_titulo.font.name = 'Calibri'

    # ── Subtítulo do bloco ────────────────────────────────────────────────────
    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub1, before=0, after=60)
    r1 = sub1.add_run("Auditoria de Dados, BI e Transparência Reguladora")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = ABAR_AZUL
    r1.font.name = 'Calibri'

    # ── Informações do curso ──────────────────────────────────────────────────
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub2, before=0, after=40)
    r2 = sub2.add_run("Curso ABAR — Infraestrutura de Gás: Medições Inteligentes e Gestão Integrada")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r2.font.name = 'Calibri'

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub3, before=0, after=80)
    r3 = sub3.add_run("Fevereiro de 2026  |  Prof. Raul Araujo")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r3.italic = True
    r3.font.name = 'Calibri'

    # Linha separadora azul
    add_horizontal_line(doc, color_hex="003886")

    # ── Campo nome do aluno ───────────────────────────────────────────────────
    nome_para = doc.add_paragraph()
    set_paragraph_spacing(nome_para, before=120, after=40)
    rn1 = nome_para.add_run("Nome do Aluno: ")
    rn1.bold = True
    rn1.font.size = Pt(11)
    rn1.font.name = 'Calibri'
    rn2 = nome_para.add_run("_" * 55)
    rn2.font.size = Pt(11)
    rn2.font.name = 'Calibri'

    data_para = doc.add_paragraph()
    set_paragraph_spacing(data_para, before=0, after=80)
    rd1 = data_para.add_run("Data: ")
    rd1.bold = True
    rd1.font.size = Pt(11)
    rd1.font.name = 'Calibri'
    rd2 = data_para.add_run("_" * 20)
    rd2.font.size = Pt(11)
    rd2.font.name = 'Calibri'

    if versao_professor:
        aviso = doc.add_paragraph()
        aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(aviso, before=0, after=80)
        shade_paragraph(aviso, "FFF2CC")
        r_aviso = aviso.add_run("  ★  VERSÃO PROFESSOR — COM GABARITO  ★  ")
        r_aviso.bold = True
        r_aviso.font.size = Pt(11)
        r_aviso.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        r_aviso.font.name = 'Calibri'

    # Linha separadora
    add_horizontal_line(doc, color_hex="003886")

    # ── Instruções ────────────────────────────────────────────────────────────
    instr = doc.add_paragraph()
    set_paragraph_spacing(instr, before=100, after=80)
    ri = instr.add_run(
        "Instruções: Leia cada questão com atenção e marque a alternativa que você considera correta. "
        "Cada questão vale 1 ponto. Total: 10 pontos."
    )
    ri.font.size = Pt(10)
    ri.italic = True
    ri.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    ri.font.name = 'Calibri'

    # ── Questões ──────────────────────────────────────────────────────────────
    for q in QUESTOES:
        num = q["num"]
        gabarito_letra = GABARITO[num]

        # Número + enunciado
        q_para = doc.add_paragraph()
        set_paragraph_spacing(q_para, before=140, after=30)

        run_num = q_para.add_run(f"Questão {num:02d}. ")
        run_num.bold = True
        run_num.font.size = Pt(11.5)
        run_num.font.color.rgb = ABAR_AZUL
        run_num.font.name = 'Calibri'

        run_enunc = q_para.add_run(q["enunciado"])
        run_enunc.bold = True
        run_enunc.font.size = Pt(11)
        run_enunc.font.name = 'Calibri'
        run_enunc.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Alternativas
        for letra, texto in q["alternativas"]:
            alt_para = doc.add_paragraph()
            set_paragraph_spacing(alt_para, before=20, after=20)
            # Indentação
            pPr = alt_para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            for existing in pPr.findall(qn('w:ind')):
                pPr.remove(existing)
            pPr.append(ind)

            is_correta = (letra == gabarito_letra)

            if versao_professor and is_correta:
                # Destaca a alternativa correta em azul e negrito
                run_letra = alt_para.add_run(f"({letra.upper()})  ")
                run_letra.bold = True
                run_letra.font.size = Pt(11)
                run_letra.font.color.rgb = COR_GABARITO
                run_letra.font.name = 'Calibri'

                run_texto = alt_para.add_run(texto + "  ✓")
                run_texto.bold = True
                run_texto.font.size = Pt(11)
                run_texto.font.color.rgb = COR_GABARITO
                run_texto.font.name = 'Calibri'
            else:
                run_letra = alt_para.add_run(f"({letra.upper()})  ")
                run_letra.font.size = Pt(11)
                run_letra.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                run_letra.font.name = 'Calibri'

                run_texto = alt_para.add_run(texto)
                run_texto.font.size = Pt(11)
                run_texto.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                run_texto.font.name = 'Calibri'

        # Campo de marcação para aluno
        if not versao_professor:
            marca_para = doc.add_paragraph()
            set_paragraph_spacing(marca_para, before=30, after=50)
            pPr = marca_para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            for existing in pPr.findall(qn('w:ind')):
                pPr.remove(existing)
            pPr.append(ind)
            rm = marca_para.add_run("Resposta: (   )")
            rm.font.size = Pt(10)
            rm.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            rm.font.name = 'Calibri'

        # Campo gabarito para professor
        if versao_professor:
            gab_para = doc.add_paragraph()
            set_paragraph_spacing(gab_para, before=30, after=50)
            shade_paragraph(gab_para, "E8F0FE")
            pPr = gab_para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            for existing in pPr.findall(qn('w:ind')):
                pPr.remove(existing)
            pPr.append(ind)
            rg1 = gab_para.add_run("Gabarito: ")
            rg1.bold = True
            rg1.font.size = Pt(10)
            rg1.font.color.rgb = COR_GABARITO
            rg1.font.name = 'Calibri'
            rg2 = gab_para.add_run(f"Alternativa ({gabarito_letra.upper()})")
            rg2.bold = True
            rg2.font.size = Pt(10)
            rg2.font.color.rgb = ABAR_VERDE
            rg2.font.name = 'Calibri'

    # ── Gabarito consolidado (só versão professor) ────────────────────────────
    if versao_professor:
        doc.add_paragraph()
        add_horizontal_line(doc, color_hex="003886")

        gab_titulo = doc.add_paragraph()
        gab_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(gab_titulo, before=120, after=60)
        rgt = gab_titulo.add_run("GABARITO CONSOLIDADO")
        rgt.bold = True
        rgt.font.size = Pt(13)
        rgt.font.color.rgb = ABAR_AZUL
        rgt.font.name = 'Calibri'

        # Tabela de gabarito
        table = doc.add_table(rows=2, cols=10)
        table.style = 'Table Grid'

        # Cabeçalho da tabela
        header_cells = table.rows[0].cells
        for i in range(10):
            cell = header_cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Q{i+1:02d}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Calibri'
            # Fundo azul ABAR
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '003886')
            for existing in tcPr.findall(qn('w:shd')):
                tcPr.remove(existing)
            tcPr.append(shd)

        # Linha de respostas
        answer_cells = table.rows[1].cells
        for i in range(10):
            q_num = i + 1
            letra = GABARITO[q_num].upper()
            cell = answer_cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(letra)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = ABAR_AZUL
            run.font.name = 'Calibri'

        # Texto consolidado abaixo da tabela
        gab_texto = doc.add_paragraph()
        set_paragraph_spacing(gab_texto, before=80, after=40)
        gab_texto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rgb = gab_texto.add_run(
            "Q1=B  |  Q2=C  |  Q3=B  |  Q4=C  |  Q5=B  |  "
            "Q6=C  |  Q7=B  |  Q8=B  |  Q9=B  |  Q10=B"
        )
        rgb.bold = True
        rgb.font.size = Pt(11)
        rgb.font.color.rgb = ABAR_AZUL
        rgb.font.name = 'Calibri'

    return doc


# ── Script principal ──────────────────────────────────────────────────────────
def main():
    print("Criando documentos de avaliação ABAR — Bloco 6...")
    print()

    for versao_prof, out_path in [(False, OUT_ALUNO), (True, OUT_PROF)]:
        label = "Professor (com gabarito)" if versao_prof else "Aluno (sem gabarito)"
        print(f"[1/3] Construindo versão {label}...")

        doc = build_document(versao_professor=versao_prof)

        print(f"[2/3] Aplicando cabeçalho e rodapé ABAR...")
        tmp_built = out_path.replace('.docx', '_built.docx')
        doc.save(tmp_built)

        # Integra header/footer do template
        final_tmp = copy_header_footer_from_template(doc, TEMPLATE_PATH)

        # Move para o caminho final
        if os.path.exists(out_path):
            os.remove(out_path)
        shutil.move(final_tmp, out_path)

        # Limpa temporários
        if os.path.exists(tmp_built):
            os.remove(tmp_built)

        print(f"[3/3] Salvo em: {out_path}")
        print()

    print("=" * 60)
    print("Documentos criados com sucesso!")
    print(f"  Versão Aluno:     {OUT_ALUNO}")
    print(f"  Versão Professor: {OUT_PROF}")
    print("=" * 60)


if __name__ == "__main__":
    main()
